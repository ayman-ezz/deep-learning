from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os

def markdown_to_docx(md_path, docx_path, image_dir):
    doc = Document()
    
    # Set up styles (basic)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    # Simple Markdown Parser
    for line in lines:
        line = line.strip()
        
        # Headers
        if line.startswith('#'):
            level = len(line.split()[0])
            text = line.lstrip('#').strip()
            # Title
            if level == 1:
                doc.add_heading(text, 0)
            else:
                doc.add_heading(text, level)
            continue
        
        # Horizontal Rule
        if line.startswith('---'):
            doc.add_paragraph("_" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
            
        # Images: ![Alt](path)
        img_match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
        if img_match:
            caption = img_match.group(1)
            img_path = img_match.group(2)
            
            # Resolve Path
            full_img_path = os.path.join(image_dir, img_path)
            if not os.path.exists(full_img_path):
                 # Try removing file:/// prefix if present
                 if img_path.startswith('file:///'):
                     clean_path = img_path.replace('file:///', '')
                     # On windows this might need drive letter fix? No, simpler to just assume filenames
                     # if the artifact dir is flat.
                     filename = os.path.basename(clean_path)
                     full_img_path = os.path.join(image_dir, filename)
            
            if os.path.exists(full_img_path):
                print(f"Adding image: {full_img_path}")
                try:
                    # Add image centered
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(full_img_path, width=Inches(6.0))
                    
                    # Add caption
                    if caption:
                         caption_p = doc.add_paragraph(caption)
                         caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                         caption_p.style = 'Caption'
                except Exception as e:
                    doc.add_paragraph(f"[Error loading image: {img_path}]")
                    print(f"Error loading image {img_path}: {e}")
            else:
                doc.add_paragraph(f"[Image not found: {img_path}]")
                print(f"Image not found: {full_img_path}")
            continue

        # Lists (Bullets)
        if line.startswith('- ') or line.startswith('* '):
            text = line[2:]
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text) # Simple bold removal for now, or keep
            # To handle bold correctly we need richer parsing, but python-docx run manipulation is verbose.
            # We'll stick to text content for simplicity, removing markdown bold syntax * symbols
            # Actually, let's just strip ** for cleanliness
            clean_text = text.replace('**', '')
            p = doc.add_paragraph(clean_text, style='List Bullet')
            # If original had bold, we lost it. 
            continue

        # Normal Text
        if line:
            # Handle Bold **text**
            # Split by **
            parts = re.split(r'(\*\*.*?\*\*)', line)
            p = doc.add_paragraph()
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
        
    doc.save(docx_path)
    print(f"Report saved to {docx_path}")

if __name__ == "__main__":
    artifact_dir = r"C:\Users\PC\.gemini\antigravity\brain\6521599c-c96e-4393-98c2-31c1f424101c"
    md_file = os.path.join(artifact_dir, "project_report.md")
    output_file = r"C:\Users\PC\Desktop\NH_LDM_Project_Report.docx"
    
    markdown_to_docx(md_file, output_file, artifact_dir)
