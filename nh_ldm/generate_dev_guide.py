
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

def generate_dev_guide():
    document = Document()

    # --- Styles ---
    style = document.styles['Normal']
    font = style.font
    font.name = 'Consolas'
    font.size = Pt(10)

    h1_style = document.styles['Heading 1']
    h1_style.font.name = 'Arial'
    h1_style.font.size = Pt(16)
    h1_style.font.color.rgb = RGBColor(0, 51, 102)

    h2_style = document.styles['Heading 2']
    h2_style.font.name = 'Arial'
    h2_style.font.size = Pt(14)
    h2_style.font.color.rgb = RGBColor(0, 102, 204)

    # --- Title Page ---
    document.add_heading('NH-LDM Developer Guide', 0)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('\n\n')
    p.add_run('Technical Reference & Onboarding Manual (Part 2)\n').bold = True
    p.add_run('Version 1.0\n\n')
    p.add_run('For Internal Development Team\n')
    document.add_page_break()

    # --- Helper ---
    def add_code_block(code):
        p = document.add_paragraph()
        run = p.add_run(code)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        paragraph_format = p.paragraph_format
        paragraph_format.left_indent = Inches(0.5)
        p.style = document.styles['No Spacing']

    # --- 1. Getting Started ---
    document.add_heading('1. Getting Started', level=1)
    document.add_paragraph(
        "Welcome to the NH-LDM project. This guide is designed to help you understand the codebase, "
        "reproduce results, and extend the functionality of the Neuro-Hybrid Latent Diffusion Model.\n\n"
        "Note: For a high-level overview of the project background and results, please refer to the 'NH-LDM Project Report' (Part 1)."
    )

    document.add_heading('1.1 Environment Setup', level=2)
    document.add_paragraph("The project relies on Python 3.8+ and PyTorch. Dependencies are listed in `requirements.txt`.")
    add_code_block(
        "# Clone the repository\n"
        "git clone https://github.com/your-org/nh-ldm.git\n\n"
        "# Install dependencies\n"
        "pip install -r requirements.txt\n"
    )

    document.add_heading('1.2 Configuration', level=2)
    document.add_paragraph(
        "All hyperparameters and paths are managed in `config.yaml`. Do not hardcode paths in scripts."
    )
    add_code_block(
        "data:\n"
        "  nsd_path: \"./data/nsd\"  # Update this to point to your local NSD copy\n"
        "diffusion:\n"
        "  device: \"cuda\"        # Change to 'cpu' for debugging on laptop\n"
    )

    # --- 2. Codebase Orientation ---
    document.add_heading('2. Codebase Orientation', level=1)
    document.add_paragraph("The project is structured as follows:")
    
    structure = (
        "project/\n"
        "├── data/            # Data loading (NSD dataset wrappers)\n"
        "│   ├── preprocessing.py\n"
        "│   └── roi_masks.py\n"
        "├── models/          # Neural network architectures\n"
        "│   ├── semantic_stream.py   # Ridge Regression for CLIP\n"
        "│   ├── structural_stream.py # Ridge Regression for VAE\n"
        "│   └── diffusion_pipeline.py\n"
        "├── train.py         # Main training loop\n"
        "└── inference.py     # Reconstruction script"
    )
    add_code_block(structure)

    # --- 3. Architecture Deep Dive ---
    document.add_heading('3. Architecture Implementation', level=1)
    
    document.add_paragraph(
        "The dual-stream architecture is implemented via two distinct decoder classes. "
        "See `project/models/semantic_stream.py` and `project/models/structural_stream.py`."
    )

    if os.path.exists('report_assets/fig1_architecture.png'):
        document.add_picture('report_assets/fig1_architecture.png', width=Inches(6))
        document.add_paragraph('Figure 1: Architecture Reference', style='Caption')

    document.add_heading('3.1 Semantic Stream (HVC -> CLIP)', level=2)
    document.add_paragraph(
        "The `SemanticDecoder` class maps voxels from the Higher Visual Cortex to CLIP text embeddings. "
        "It uses a multi-output Ridge Regression. We select the top N voxels (default 1000) based on correlation during training."
    )
    add_code_block(
        "class SemanticDecoder:\n"
        "    def train(self, X, Y):\n"
        "        # voxel selection logic...\n"
        "        self.voxel_indices = select_top_voxels(X, Y, keep=self.n_voxels)\n"
        "        # fit ridge...\n"
    )

    document.add_heading('3.2 Structural Stream (EVC -> VAE)', level=2)
    document.add_paragraph(
        "The `StructuralDecoder` maps Early Visual Cortex voxels to the 64x64x4 latent space of the Stable Diffusion VAE. "
        "Since the dimensionality is high (16384 outputs), we use a more aggressive regularization or dimensionality reduction if needed."
    )

    # --- 4. Data Pipeline ---
    document.add_heading('4. Data Pipeline', level=1)
    document.add_paragraph(
        "Data loading is handled by `project.data.preprocessing.DataProcessor`. "
        "This class abstracts away the complexity of HDF5 files from the NSD dataset."
    )
    document.add_paragraph("Key responsibilities:", style='List Bullet')
    document.add_paragraph("Loading fMRI betas", style='List Bullet')
    document.add_paragraph("Applying ROI masks (using `roi_masks.py`)", style='List Bullet')
    document.add_paragraph("Z-score normalization (statistics computed on training set)", style='List Bullet')

    if os.path.exists('report_assets/fig2_roi_distribution.png'):
        document.add_picture('report_assets/fig2_roi_distribution.png', width=Inches(5))
        document.add_paragraph('Reference: Voxel Counts per ROI', style='Caption')

    # --- 5. Training & Inference ---
    document.add_heading('5. Workflow guide', level=1)
    
    document.add_heading('5.1 Training (`train.py`)', level=2)
    document.add_paragraph(
        "The training script `train.py` orchestrates the fitting of both decoders. "
        "It saves the trained models as `.pkl` files in the `project/models/` directory."
    )
    document.add_paragraph("To start training:")
    add_code_block("python -m project.train --config config.yaml")

    document.add_heading('5.2 Inference (`inference.py`)', level=2)
    document.add_paragraph(
        "Inference loads the pickled models and the Stable Diffusion pipeline. "
        "It iterates through the test set, predicts latents/embeddings, and runs the diffusion loop."
    )
    document.add_paragraph("To run inference:")
    add_code_block("python -m project.inference --output_dir results/")

    # --- 6. Troubleshooting ---
    document.add_heading('6. Troubleshooting', level=1)
    
    table = document.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Issue'
    hdr_cells[1].text = 'Solution'
    
    row = table.add_row().cells
    row[0].text = "CUDA Out of Memory"
    row[1].text = "Reduce `batch_size` in config.yaml or set `device: cpu`."
    
    row = table.add_row().cells
    row[0].text = "FileNotFoundError: nsd_stim_info.pkl"
    row[1].text = "Ensure you have downloaded the NSD metadata and placed it in the data root."

    # --- 7. Extending the Project ---
    document.add_heading('7. Extending the Project', level=1)
    document.add_paragraph(
        "To add a new subject, update the `subjects` list in `config.yaml`. "
        "Note that you will need to re-run training as voxel mappings are subject-specific."
    )

    output_path = 'NH_LDM_Developer_Guide.docx'
    document.save(output_path)
    print(f"Developer Guide generated: {output_path}")

if __name__ == "__main__":
    generate_dev_guide()
